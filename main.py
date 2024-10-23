from docx import Document
from docx.shared import Pt
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import tkinter as tk
from tkinter import messagebox
from docx.shared import Pt
import locale
from docx.enum.section import WD_SECTION
from docx.shared import Cm
from datetime import datetime
from interface import abrir_interface
import os



# Função para criar o documento com as informações fornecidas
def gerar_documento(concessionaria, estacao, tipo_local, data, horário, imagens, legendas):
    document = Document()

    # Configurações da página
    section = document.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(1)
    section.bottom_margin = Cm(1.5)

    # Definir fonte padrão
    style = document.styles['Normal']
    font = style.font
    font.size = Pt(12)
    font.name = 'Palatino Linotype'

    # Função para adicionar borda à página
    def add_page_border(section):
        sectPr = section._sectPr
        pgBorders = OxmlElement('w:pgBorders')
        for border_attr in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_attr}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '6')
            border.set(qn('w:color'), '000000')
            border.set(qn('w:space'), '50')
            pgBorders.append(border)
        sectPr.append(pgBorders)

    add_page_border(section)

    preposicao = "no" if tipo_local == "Terminal" else "na"

    # Cabeçalho
    header = section.header
    table = header.add_table(rows=1, cols=3, width=Cm(35))
    table.columns[0].width = Cm(5)
    table.columns[1].width = Cm(20)
    table.columns[2].width = Cm(5)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER

    cell_logo = table.cell(0, 0)
    run_logo = cell_logo.paragraphs[0].add_run()
    # Adicione o caminho da imagem aqui
    run_logo.add_picture('Logo AGETRANSP.png', width=Inches(1))

    cell_text = table.cell(0, 1)
    text_run = cell_text.paragraphs[0].add_run("AGÊNCIA REGULADORA DE SERVIÇOS PÚBLICOS CONCEDIDOS DE TRANSPORTES AQUAVIÁRIOS, "
                                               "FERROVIÁRIOS E METROVIÁRIOS E DE RODOVIAS DO ESTADO DO RIO DE JANEIRO.")
    text_run.font.size = Pt(7)
    text_run.bold = True
    cell_text.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    cell_page = table.cell(0, 2)
    cell_page.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    cell_page.paragraphs[0].add_run("Página ").font.size = Pt(8)

    # Título da página de capa
    capa = document.add_paragraph(f'RELATÓRIO TÉCNICO \n xxx/CATRA/2024')
    capa.alignment = WD_ALIGN_PARAGRAPH.CENTER
    capa.runs[0].font.size = Pt(24)
    capa.runs[0].bold = True

    document.add_paragraph('')
    paragrafo_1 = document.add_paragraph(f'Atividades de Fiscalização {preposicao} {tipo_local} {estacao}')
    paragrafo_1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragrafo_1.runs[0].font.size = Pt(18)
    paragrafo_1.runs[0].bold = True

    document.add_paragraph('')
    paragrafo_2 = document.add_paragraph(f'Processo SEI-100003/00xxxx/2024')
    paragrafo_2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragrafo_2.runs[0].font.size = Pt(18)

    document.add_paragraph('')
    paragrafo_3 = document.add_paragraph(f'Concessionária {concessionaria}')
    paragrafo_3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragrafo_3.runs[0].font.size = Pt(20)

    document.add_paragraph('')
    document.add_paragraph('')
    paragrafo_4 = document.add_paragraph('\nElaboração\nCATRA – Câmara de Transportes e Rodovias')
    paragrafo_4.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragrafo_4.runs[0].font.size = Pt(14)

    document.add_paragraph('')
    document.add_paragraph('')
    document.add_paragraph('')
    locale.setlocale(locale.LC_TIME, 'pt_BR.UTF-8') 
    mes_atual = datetime.now().strftime('%B').capitalize()  # Nome do mês atual
    ano_atual = datetime.now().year  # Ano atual
    paragrafo_5 = document.add_paragraph(mes_atual + ' de ' + str(ano_atual))
    run5 = paragrafo_5.runs[0]
    run5.italic = True
    run5.bold = True
    run5.font.size = Pt(14)
    paragrafo_5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_page_break()

    # Índice
    document.add_paragraph('')  # Parágrafo vazio
    indice = document.add_paragraph('ÍNDICE')
    indice1 = indice.runs[0]
    indice1.bold = True
    indice1.font.size = Pt(12)

    indice.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Quebra de página
    document.add_page_break()


    heading_intro = document.add_heading('1. INTRODUÇÃO', level=1)
    run_intro = heading_intro.runs[0]
    run_intro.font.name = 'Palatino Linotype'
    run_intro.font.color.rgb = RGBColor(0, 51, 102)  # Definindo a cor RGB

    document.add_paragraph(
        f'O presente Relatório Técnico apresenta os resultados da vistoria realizada em {data}, '
        f'{preposicao} {tipo_local} de {estacao}, sob a operação da Concessionária {concessionaria}.'
    )# Desenvolvimento
    heading_dev = document.add_heading('2. DESENVOLVIMENTO', level=1)
    run_dev = heading_dev.runs[0]
    run_dev.font.name = 'Palatino Linotype'
    run_dev.font.color.rgb = RGBColor(0, 51, 102)  # Definindo a cor RGB

    document.add_paragraph(
        f'No dia {data}, por volta das {horário} Câmara de Transportes e Rodovias da AGETRANSP realizaram atividade de inspeção técnica {preposicao} {tipo_local} de {estacao}. O objetivo foi verificar a acessibilidade, condições de '
        'operacionalidade, limpeza e conservação, além da comunicação sonora e visual.'
    ).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  

    # Atividades Realizadas
    heading_atividades = document.add_heading('2.1. Atividades Realizadas', level=2)
    run_atividades = heading_atividades.runs[0]
    run_atividades.font.name = 'Palatino Linotype'
    run_atividades.font.color.rgb = RGBColor(0, 51, 102)  # Definindo a cor RGB

    document.add_paragraph(
        'Durante a atividade de fiscalização foram avaliados aspectos como acessibilidade, limpeza e conservação das plataformas, intervalos, mezaninos e assentos, itens de comunicação visual, conforme registros fotográficos apresentados a seguir:'
    ).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  


    num_columns = 2
    table = document.add_table(rows=1, cols=num_columns)  # Cria uma tabela com 2 colunas
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Variável para acompanhar a linha atual da tabela
    current_row = table.add_row()

    for i, image_path in enumerate(imagens):
        # Se a coluna for preenchida, adicionar uma nova linha
        if i % num_columns == 0 and i != 0:
            current_row = table.add_row()
        
        # Adicionar a imagem na célula correspondente
        cell = current_row.cells[i % num_columns]
        
        # Remover todos os parágrafos da célula para ter controle total sobre o conteúdo
        cell._element.clear_content()


        # Adicionar parágrafo para a imagem
        paragraph_image = cell.add_paragraph()
        run_image = paragraph_image.add_run()
        run_image.add_picture(image_path, width=Cm(6.4), height=Cm(6.4))

        # Centralizar a imagem
        paragraph_image.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Adicionar legenda logo abaixo da imagem, sem criar novo parágrafo
        run_image.add_break()
        legenda_atual = legendas[i] if i < len(legendas) else f'Figura {i+1}: INSERIR LEGENDA' 
        run_caption = paragraph_image.add_run(f'Figura {i+1}: {legenda_atual}')
        run_caption.font.size = Pt(7)  
        
        paragraph_image.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Ajustar espaçamento (aqui garante que não haja espaço extra)
        paragraph_image.paragraph_format.space_before = Pt(0)
        paragraph_image.paragraph_format.space_after = Pt(2)


    # Constatações
    heading_constatacoes = document.add_heading('3. CONSTATAÇÕES', level=1)
    run_constatacoes = heading_constatacoes.runs[0]
    run_constatacoes.font.name = 'Palatino Linotype'
    run_constatacoes.font.color.rgb = RGBColor(0, 51, 102)  # Definindo a cor RGB

    document.add_paragraph(
        'Durante as vistorias, foram verificados aspectos fundamentais para uma boa experiência dos usuários, incluindo condições de conservação, limpeza, manutenção das faixas amarelas de segurança e piso tátil, e operacionalidade dos equipamentos instalados nas estações. As observações realizadas durante a inspeção são as seguintes:'
    ).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  

    
    heading_local = document.add_heading(f'3.1. {estacao}', level=2)
    run_local = heading_local.runs[0]
    run_local.font.name = 'Palatino Linotype'
    run_local.font.color.rgb = RGBColor(0, 51, 102)  # Definindo a cor RGB

    document.add_paragraph(
        'Limpeza e conservação: A limpeza geral foi considerada adequada, no entanto, foi observado acúmulo de lixo próximo à plataforma 2 e '
        'rachaduras na cobertura da plataforma.'
    ).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  


    document.add_paragraph(
        'Comunicação sonora e visual: A comunicação sonora estava em funcionamento, mas os painéis eletrônicos estavam inoperantes, prejudicando a comunicação visual.'
    ).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  

    document.add_paragraph(
        'Acessibilidade: Foi constatada a ausência de piso tátil em diversas áreas da estação, além de desnível entre o trem e a plataforma, o que '
        'não atende às normas técnicas de acessibilidade.'
    ).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  

    # Conclusão
    heading_conclusao = document.add_heading('4. CONCLUSÃO', level=1)
    run_conclusao = heading_conclusao.runs[0]
    run_conclusao.font.name = 'Palatino Linotype'
    run_conclusao.font.color.rgb = RGBColor(0, 51, 102)  # Definindo a cor RGB

    document.add_paragraph(
        f'Diante das constatações apresentadas, recomenda-se que a concessionária {concessionaria} realize os reparos necessários e '
        'implementação das medidas corretivas no prazo máximo de 30 dias, a fim de garantir a segurança e acessibilidade.'
    ).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  


    # Assinaturas
    data_atual = datetime.now().strftime('%d de %B de %Y') 
    document.add_paragraph('')
    document.add_paragraph(f'Em {data_atual},\n')

    #document.add_paragraph('Isaque da Cunha Soares\nEstagiário da Câmara de Transportes e Rodovias')
    #document.add_paragraph('João Gabriel Lopes Zarur\nAgente de Fiscalização da Câmara de Transportes e Rodovias')
    #document.add_paragraph('Rafael Lanunci da Silva Teixeira Poubel\nGerente da Câmara de Transportes e Rodovias')

    # Salvando o documento
    document.save(f'Relatorio_Fiscalizacao_{estacao}.docx')

abrir_interface(gerar_documento)
